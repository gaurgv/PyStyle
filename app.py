from flask import Flask, request, render_template, send_file
from docx import Document
import os
import re

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    doc_file = request.files['doc']
    template_file = request.files['template']

    doc_path = os.path.join(app.config['UPLOAD_FOLDER'], doc_file.filename)
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file.filename)
    
    doc_file.save(doc_path)
    template_file.save(template_path)
    
    # Process the chapter
    styled_doc_path = process_document(doc_path, template_path)
    
    return send_file(styled_doc_path, as_attachment=True)

def process_document(doc_path, template_path):
    doc = Document(doc_path)
    #keywords_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'keyword_list.txt')
    book_id = os.path.basename(doc_path).split('_')[0]
    chapter_number = doc.paragraphs[0].text.strip()
    # Initialize a counter for figure numbers
    figure_counter = 1
    # Regular expression to match figure captions
    figure_caption_pattern = re.compile(rf"^Figure {chapter_number}\.\d+:")
    # Regular expression to identify code in text
    code_term_pattern = re.compile(r'\b\w+\(\)|\b[A-Za-z][a-z]+[A-Z]\w*\b|\b\w+_\w+\b')

    # Heading styles from the template
    heading_styles = {
        'Heading 1': 'HS - Heading_1 [PACKT]',
        'Heading 2': 'HS - Heading_2 [PACKT]',
        'Heading 3': 'HS - Heading_3 [PACKT]',
        'Heading 4': 'HS - Heading_4 [PACKT]',
        'Heading 5': 'HS - Heading_5 [PACKT]',
        'Heading 6': 'HS - Heading_6 [PACKT]',
    }

    # Iterate through paragraphs and apply heading styles
    for para in doc.paragraphs:
        # Check if the paragraph has a heading style
        if para.style.name in heading_styles:
            # Apply the corresponding style
            if para.style.name not in heading_styles.values():
                para.style = heading_styles[para.style.name]
    
    if len(doc.paragraphs) >= 2:
        chapter_number_paragraph = doc.paragraphs[0]
        chapter_title_paragraph = doc.paragraphs[1]

        # Apply styles for chapter number and title
        chapter_number_paragraph.style = 'HS - ChapterNumber [PACKT]'
        chapter_title_paragraph.style = 'HS - ChapterTitle [PACKT]'

    
    # First, apply the Normal text style to paragraphs
    for paragraph in doc.paragraphs:
        # Skip chapter number, title, and any headings
        if paragraph == doc.paragraphs[0] or paragraph == doc.paragraphs[1] or paragraph.style.name.startswith("HS"):
            continue
        
        # Skip figure captions and layout information
        if figure_caption_pattern.match(paragraph.text):
            continue

        # Apply Normal text style to regular paragraphs that are not indented or lists
        if not paragraph.paragraph_format.left_indent and not paragraph.style.name.startswith("List"):
            paragraph.style = 'P0 - Normal [PACKT]'
            # Apply ""CS - InlineCode [PACKT]" style to identified code within texts
            for match in code_term_pattern.finditer(paragraph.text):
                start, end = match.span()
                run = paragraph.add_run(paragraph.text[start:end])
                run.style = 'CS - InlineCode [PACKT]'
                paragraph.text = paragraph.text[:start] + paragraph.text[end:]

    # Handle adding layout information after figures
    for i, paragraph in enumerate(doc.paragraphs):
        if figure_caption_pattern.match(paragraph.text):
            # layout information string
            layout_info = f"Insert image {book_id}_{chapter_number.zfill(2)}_{str(figure_counter).zfill(2)}"

            # A new paragraph for the layout information
            layout_info_paragraph = doc.add_paragraph(layout_info)
            layout_info_paragraph.style = 'PF - LayoutInformation [PACKT]'

            # Move the new paragraph after the figure caption
            p = paragraph._element
            p.addnext(layout_info_paragraph._element)

            # Increment the figure counter
            figure_counter += 1

    #identify_and_style_keywords(doc, keywords_file_path)
    identify_and_style_urls(doc)
    styled_doc_path = os.path.join(app.config['UPLOAD_FOLDER'], 'styled_' + os.path.basename(doc_path))
    doc.save(styled_doc_path)
    return styled_doc_path

# def identify_and_style_keywords(doc, keywords_file_path):
#     # Load keywords
#     with open(keywords_file_path, 'r') as f:
#         keywords_set = set(line.strip() for line in f if line.strip())

#     keyword_style = "CS - KeyWord [PACKT]"

#     # Regex pattern for keywords
#     keyword_pattern = re.compile(r'\b([A-Za-z\s]+)\s\(([A-Z]+)\)')

#     for para in doc.paragraphs:
#         matches = keyword_pattern.finditer(para.text)
#         if matches:
#             # Split paragraph into runs to apply styles selectively
#             original_text = para.text
#             para.clear()  # Clear the paragraph's existing runs
#             cursor = 0

#             for match in matches:
#                 full_term_start, full_term_end = match.span(1)  # Span of the full term
#                 abbr_start, abbr_end = match.span(2)  # Span of the abbreviation

#                 # Combine the full term and abbreviation
#                 keyword = f"{match.group(1)} ({match.group(2)})"

#                 # Add text before the match
#                 if cursor < full_term_start:
#                     para.add_run(original_text[cursor:full_term_start])

#                 # Apply the style
#                 if keyword in keywords_set:
#                     full_term_run = para.add_run(original_text[full_term_start:full_term_end])
#                     full_term_run.style = keyword_style
#                     abbr_run = para.add_run(original_text[abbr_start - 1:abbr_end + 1])  # Include parentheses
#                     abbr_run.style = keyword_style
#                 else:
#                     # Normal text style if no match
#                     para.add_run(original_text[full_term_start:abbr_end + 1])

#                 cursor = abbr_end + 1

#             # Add remaining text after the last match
#             if cursor < len(original_text):
#                 para.add_run(original_text[cursor:])

def identify_and_style_urls(doc):
    url_style = "CS - URL [PACKT]"
    url_pattern = re.compile(r'(?:\(|\[)?(https?://(?:www\.)?[^\s]+)(?:\)|\])?')
    for para in doc.paragraphs:
        # Skip paragraphs that contain images
        if "graphic" in para._p.xml:
            continue

        # Apply the URL style
        for run in para.runs:
            if run.style and run.style.name == "Hyperlink":
                run.style = url_style

        # Identify and style URLs using regex
        matches = url_pattern.finditer(para.text)
        if matches:
            # Reconstruct the paragraph to preserve existing non-text elements
            original_text = para.text
            runs = list(para.runs)  # Backup runs for non-text content
            para.clear()  # Clear the paragraph

            cursor = 0
            for match in matches:
                full_span = match.span()  # Span of the entire match (including brackets/parentheses)
                url_span = match.span(1)  # Span of the captured URL only

                # Add text before the match
                if cursor < full_span[0]:
                    para.add_run(original_text[cursor:full_span[0]])

                if full_span[0] < url_span[0]:
                    para.add_run(original_text[full_span[0]:url_span[0]])

                # Add the URL with the URL style
                url_run = para.add_run(original_text[url_span[0]:url_span[1]])
                url_run.style = url_style

                if url_span[1] < full_span[1]:
                    para.add_run(original_text[url_span[1]:full_span[1]])

                cursor = full_span[1]

            # Add remaining text after the last match
            if cursor < len(original_text):
                para.add_run(original_text[cursor:])

            # Readd images
            for run in runs:
                if run._element.tag.endswith("drawing") or run._element.tag.endswith("object"):
                    para._p.append(run._element)

if __name__ == '__main__':
    app.run(debug=True)